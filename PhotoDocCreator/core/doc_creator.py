from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import logging
from PIL import Image
import tempfile

logger = logging.getLogger(__name__)

class DocumentCreator:
    def __init__(self, config):
        self.config = config
        self.doc = None
        self.temp_files = []  # Для хранения временных файлов
    
    def __del__(self):
        """Очистка временных файлов при удалении объекта"""
        self.cleanup_temp_files()
    
    def cleanup_temp_files(self):
        """Удаляет все временные файлы"""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except Exception as e:
                logger.warning(f"Не удалось удалить временный файл {temp_file}: {e}")
        self.temp_files = []
    
    def create_document(self, image_data_list, log_callback=None):
        """Создает документ Word с фотографиями"""
        try:
            if log_callback:
                log_callback("🚀 Начало создания документа...")
            
            # Создаем документ
            self.doc = Document()
            self._setup_page_layout()
            
            # Добавляем заголовки
            self._add_titles()
            
            # Добавляем фотографии
            added_count = self._add_images(image_data_list, log_callback)
            
            # Добавляем колонтитулы
            if self.config.get('enable_footer', True):
                self._add_footers()
                if log_callback:
                    log_callback("✅ Колонтитул добавлен")
            else:
                if log_callback:
                    log_callback("ℹ️ Колонтитул отключен")
            
            # Сохраняем документ
            output_file = self.config.get('word_file', 'output.docx')
            self._ensure_directory_exists(output_file)
            self.doc.save(output_file)
            
            # Очищаем временные файлы
            self.cleanup_temp_files()
            
            if log_callback:
                log_callback(f"✅ Готово! Создан документ с {added_count} фотографиями")
                log_callback(f"📁 Файл: {output_file}")
            
            return True, output_file, added_count
            
        except Exception as e:
            error_msg = f"❌ Критическая ошибка создания документа: {str(e)}"
            if log_callback:
                log_callback(error_msg)
            logger.error(error_msg, exc_info=True)
            self.cleanup_temp_files()
            return False, str(e), 0
    
    def _setup_page_layout(self):
        """Настраивает параметры страницы"""
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    def _add_titles(self):
        """Добавляет заголовки документа"""
        # Основной заголовок
        title_paragraph = self.doc.add_paragraph()
        title_run = title_paragraph.add_run(self.config.get('department_name', ''))
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.name = "Times New Roman"
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.paragraph_format.space_after = Pt(12)
        
        # Заголовок фототаблицы
        table_title_paragraph = self.doc.add_paragraph()
        table_title_text = self.config.get('photo_table_title', '')
        table_title_run = table_title_paragraph.add_run(table_title_text)
        table_title_run.font.size = Pt(12)
        table_title_run.font.bold = True
        table_title_run.font.name = "Times New Roman"
        table_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table_title_paragraph.paragraph_format.space_after = Pt(24)
    
    def _add_images(self, image_data_list, log_callback=None):
        """Добавляет изображения в документ"""
        images_per_page = self.config.get('images_per_page', 2)
        image_width = self.config.get('image_width', 6.0)
        image_height = self.config.get('image_height', 9.0)
        font_family = self.config.get('font_family', 'Times New Roman')
        font_size = self.config.get('font_size', 12)
        font_bold = self.config.get('font_bold', False)
        multi_folder_mode = self.config.get('multi_folder_mode', False)
        rotation_info = self.config.get('rotation_info', {})
        
        added_count = 0
        
        for i in range(0, len(image_data_list), images_per_page):
            # Добавляем разрыв страницы (кроме первой)
            if i > 0:
                self.doc.add_page_break()
                for _ in range(2):
                    self.doc.add_paragraph()
            
            # Добавляем фото на текущую страницу
            photos_on_this_page = 0
            for j in range(images_per_page):
                img_index = i + j
                if img_index >= len(image_data_list):
                    break
                    
                photo_info = image_data_list[img_index]
                img_path = photo_info['path']
                filename = photo_info.get('filename', 'Unknown')
                
                # Получаем информацию о повороте
                rotation = photo_info.get('rotation', 0)
                if not rotation and img_path in rotation_info:
                    rotation = rotation_info[img_path]
                
                # Получаем подпись
                if multi_folder_mode:
                    caption = self._get_caption_multi(photo_info)
                else:
                    caption = self._get_caption_single(photo_info)
                
                # Пытаемся добавить изображение
                success = self._add_single_image(
                    img_path, filename, image_width, image_height, 
                    caption, font_family, font_size, font_bold, 
                    log_callback, rotation
                )
                
                if success:
                    added_count += 1
                    photos_on_this_page += 1
                    
                    if log_callback:
                        log_callback(f"✅ Добавлено: {caption}")
                else:
                    if log_callback:
                        log_callback(f"❌ Не удалось добавить: {filename}")
        
        return added_count
    
    def _convert_image_for_docx(self, image_path):
        """
        Конвертирует изображение в формат, совместимый с Word
        Возвращает путь к временному файлу
        """
        try:
            with Image.open(image_path) as img:
                # Конвертируем в RGB если нужно
                if img.mode in ('RGBA', 'P', 'LA', 'CMYK'):
                    img = img.convert('RGB')
                
                # Создаем временный файл
                fd, temp_path = tempfile.mkstemp(suffix='.jpg')
                os.close(fd)
                
                # Сохраняем в JPEG с оптимальным качеством
                img.save(temp_path, 'JPEG', quality=95, optimize=True)
                self.temp_files.append(temp_path)
                
                logger.debug(f"Изображение конвертировано: {image_path} -> {temp_path}")
                return temp_path
                
        except Exception as e:
            logger.error(f"Ошибка конвертации {image_path}: {e}")
            return None
    
    def _add_single_image(self, img_path, filename, width, height, caption, font_family, font_size, font_bold, log_callback=None, rotation=0):
        """Добавляет одно изображение с подписью и поворотом"""
        try:
            # Проверяем существование файла
            if not os.path.exists(img_path):
                if log_callback:
                    log_callback(f"❌ Файл не найден: {filename}")
                return False
            
            # Проверяем, что файл является валидным изображением
            try:
                with Image.open(img_path) as img:
                    img.verify()
            except Exception as e:
                if log_callback:
                    log_callback(f"❌ Файл поврежден или не является изображением {filename}: {str(e)}")
                return False
            
            # Обработка поворота
            final_img_path = img_path
            temp_rotated_path = None
            
            if rotation != 0:
                try:
                    with Image.open(img_path) as img:
                        rotated_img = img.rotate(rotation, expand=True)
                        
                        # Создаем временный файл для повернутого изображения
                        fd, temp_rotated_path = tempfile.mkstemp(suffix='.jpg')
                        os.close(fd)
                        rotated_img.save(temp_rotated_path, 'JPEG', quality=95)
                        final_img_path = temp_rotated_path
                        self.temp_files.append(temp_rotated_path)
                        
                        if log_callback:
                            log_callback(f"↷ Изображение повернуто на {rotation}°: {filename}")
                except Exception as e:
                    if log_callback:
                        log_callback(f"⚠️ Ошибка поворота {filename}: {str(e)}")
            
            # Добавляем изображение в документ
            p_image = self.doc.add_paragraph()
            p_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_image = p_image.add_run()
            
            # Используем повернутое изображение если есть
            success = False
            try:
                run_image.add_picture(final_img_path, width=Cm(width), height=Cm(height))
                success = True
            except Exception as e1:
                logger.warning(f"Прямое добавление не удалось для {filename}: {e1}")
                
                # Пробуем через конвертацию
                try:
                    temp_path = self._convert_image_for_docx(final_img_path)
                    if temp_path and os.path.exists(temp_path):
                        run_image.add_picture(temp_path, width=Cm(width), height=Cm(height))
                        success = True
                        if log_callback:
                            logger.debug(f"Изображение {filename} добавлено через конвертацию")
                    else:
                        if log_callback:
                            log_callback(f"❌ Не удалось конвертировать изображение {filename}")
                except Exception as e2:
                    if log_callback:
                        log_callback(f"❌ Ошибка при добавлении конвертированного изображения {filename}: {str(e2)}")
            
            # Добавляем подпись если изображение было успешно добавлено
            if success:
                p_caption = self.doc.add_paragraph()
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_caption = p_caption.add_run(caption)
                
                # Настраиваем шрифт подписи
                run_caption.font.size = Pt(font_size)
                run_caption.font.name = font_family
                if font_bold:
                    run_caption.font.bold = True
                
                # Добавляем отступ
                self.doc.add_paragraph()
                
                return True
            
            return False
            
        except Exception as e:
            if log_callback:
                log_callback(f"❌ Неожиданная ошибка при добавлении {filename}: {str(e)}")
            logger.error(f"Ошибка добавления изображения {filename}: {e}", exc_info=True)
            return False
    
    def _get_caption_single(self, photo_info):
        """Генерирует подпись для одиночного режима"""
        photo_number = photo_info['global_number']
        caption_rules = self.config.get('caption_rules', [])
        
        # Проверяем правила подписей
        for rule in caption_rules:
            if len(rule) >= 3:
                start, end, text = rule[0], rule[1], rule[2]
                if start <= photo_number <= end:
                    return f"Фото № {photo_number}. {text}"
        
        # Если правил нет - стандартная подпись
        return f"Фото № {photo_number}"
    
    def _get_caption_multi(self, photo_info):
        """Генерирует подпись для многопапкового режима"""
        photo_number = photo_info['global_number']
        folder_rules = photo_info.get('folder_rules', [])
        folder_start = photo_info.get('folder_start_number', 1)
        
        # Локальный номер фото в папке
        local_photo_number = photo_number - folder_start + 1
        
        # Сначала проверяем правила папки
        for rule in folder_rules:
            if len(rule) >= 3:
                start, end, text = rule[0], rule[1], rule[2]
                if start <= local_photo_number <= end:
                    return f"Фото № {photo_number}. {text}"
        
        # Затем общие правила
        caption_rules = self.config.get('caption_rules', [])
        for rule in caption_rules:
            if len(rule) >= 3:
                start, end, text = rule[0], rule[1], rule[2]
                if start <= photo_number <= end:
                    return f"Фото № {photo_number}. {text}"
        
        return f"Фото № {photo_number}"
    
    def _add_footers(self):
        """Добавляет колонтитулы"""
        footer_text = self._generate_footer_text()
        
        for section in self.doc.sections:
            footer = section.footer
            footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_paragraph.clear()
            
            run = footer_paragraph.add_run(footer_text)
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _generate_footer_text(self):
        """Генерирует текст для колонтитула"""
        position = self.config.get('officer_position', '')
        department = self.config.get('footer_department', '')
        rank = self.config.get('officer_rank', '')
        name = self.config.get('officer_name', '')
        
        return f"{position} {department}\n{rank}\t\t\t\t\t\t{name}"
    
    def _ensure_directory_exists(self, file_path):
        """Создает директорию для файла если она не существует"""
        directory = os.path.dirname(file_path)
        if directory and not os.path.exists(directory):
            os.makedirs(directory, exist_ok=True)